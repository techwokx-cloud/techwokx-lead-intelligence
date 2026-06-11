import streamlit as st
import pandas as pd
import json
import os
import requests
import smtplib
import imaplib
import email
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from datetime import datetime, timedelta
import time
import qrcode
from io import BytesIO
import base64
import re
import tldextract
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============ PAGE CONFIG ============
st.set_page_config(
    page_title="TechWokx Lead Intelligence",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============ SESSION STATE INIT ============
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'user' not in st.session_state:
    st.session_state.user = None
if 'current_page' not in st.session_state:
    st.session_state.current_page = 'login'
if 'leads' not in st.session_state:
    st.session_state.leads = []
if 'email_log' not in st.session_state:
    st.session_state.email_log = []
if 'bounced_emails' not in st.session_state:
    st.session_state.bounced_emails = []
if 'batch_results' not in st.session_state:
    st.session_state.batch_results = []
if 'letter_queue' not in st.session_state:
    st.session_state.letter_queue = []
if 'call_queue' not in st.session_state:
    st.session_state.call_queue = []
if 'visit_queue' not in st.session_state:
    st.session_state.visit_queue = []
if 'selected_company' not in st.session_state:
    st.session_state.selected_company = None
if 'searched_companies' not in st.session_state:
    st.session_state.searched_companies = set()

# ============ FILE STORAGE ============
os.makedirs("data", exist_ok=True)
LEADS_FILE = "data/leads.json"
EMAIL_LOG_FILE = "data/email_log.json"
BOUNCE_LOG_FILE = "data/bounce_log.json"

def load_leads():
    try:
        if os.path.exists(LEADS_FILE):
            with open(LEADS_FILE, 'r') as f:
                return json.load(f)
    except:
        pass
    return []

def save_leads():
    try:
        with open(LEADS_FILE, 'w') as f:
            json.dump(st.session_state.leads, f, indent=2, default=str)
    except:
        pass

def load_email_log():
    try:
        if os.path.exists(EMAIL_LOG_FILE):
            with open(EMAIL_LOG_FILE, 'r') as f:
                return json.load(f)
    except:
        pass
    return []

def save_email_log():
    try:
        with open(EMAIL_LOG_FILE, 'w') as f:
            json.dump(st.session_state.email_log, f, indent=2, default=str)
    except:
        pass

st.session_state.leads = load_leads()
st.session_state.email_log = load_email_log()

# ============ LOGIN CREDENTIALS ============
USERS = {
    "hello@techwokx.online": {
        "password": "Gtech.5628!@#$",
        "role": "admin",
        "name": "TechWokx Admin"
    }
}

# ============ SMTP CONFIGURATION ============
SMTP_CONFIG = {
    "host": "smtp.hostinger.com",
    "port": 465,
    "username": "hello@techwokx.online",
    "password": "Gtech.5628!@#$",
    "use_ssl": True,
    "daily_limit": 50,
    "sent_today": 0
}

# ============ IMAP CONFIGURATION FOR BOUNCE DETECTION ============
IMAP_CONFIG = {
    "host": "imap.hostinger.com",
    "port": 993,
    "username": "hello@techwokx.online",
    "password": "Gtech.5628!@#$",
    "use_ssl": True
}

def check_bounced_emails():
    """Check for bounced email notifications"""
    try:
        import imaplib
        mail = imaplib.IMAP4_SSL(IMAP_CONFIG['host'], IMAP_CONFIG['port'])
        mail.login(IMAP_CONFIG['username'], IMAP_CONFIG['password'])
        mail.select('INBOX')
        
        # Search for bounce emails
        result, data = mail.search(None, '(SUBJECT "Undelivered" UNSEEN)')
        if result == 'OK':
            for num in data[0].split():
                result, msg_data = mail.fetch(num, '(RFC822)')
                if result == 'OK':
                    email_body = msg_data[0][1].decode('utf-8')
                    # Extract bounced email address
                    import re
                    bounced_match = re.search(r'<([^>]+@[^>]+)>', email_body)
                    if bounced_match:
                        bounced_email = bounced_match.group(1)
                        st.session_state.bounced_emails.append({
                            "email": bounced_email,
                            "date": datetime.now().isoformat(),
                            "reason": "Email bounced - invalid address"
                        })
                mail.store(num, '+FLAGS', '\\Seen')
        mail.close()
        mail.logout()
    except Exception as e:
        pass

# ============ COMPLETE GHANA CITIES DATABASE ============
ALL_CITIES = {
    "Greater Accra": [
        "Accra Central", "Airport Residential", "Airport City", "Cantonments", "Labone", 
        "Osu", "Ring Road Central", "Ridge", "North Ridge", "Adabraka", "Asylum Down", 
        "Tudu", "Jamestown", "Usshertown", "Korle Bu", "Mamprobi", "Chorkor", "Kaneshie", 
        "Achimota", "Legon", "Madina", "Adenta", "East Legon", "West Legon", "Dzorwulu", 
        "Shiashie", "Trasacco Valley", "Spintex", "Sakumono", "Teshie", "Nungua", "Lashibi", 
        "Ashaiman", "Tema", "Tema Community 1", "Tema Community 2", "Tema Community 3", 
        "Tema Community 4", "Tema Community 5", "Tema Community 6", "Tema Community 7", 
        "Tema Community 8", "Tema Community 9", "Tema Community 10", "Tema Industrial Area", 
        "Kpone", "Adjei Kojo", "Afienya", "Dawhenya", "Prampram", "Dodowa", "Abokobi", 
        "Danfa", "Teiman", "Taifa", "Kwabenya", "Dome", "Pokuase", "Haasto", "Amasaman", 
        "Nsawam", "Medie", "Sarpeiman", "Weija", "Gbawe", "Mallam", "Kwashieman", "Darkuman", 
        "Kasoa", "Budumburam", "Opeikuma"
    ],
    "Ashanti": [
        "Kumasi Central", "Adum", "Bantama", "Asafo", "Asokwa", "Tafo", "Suame", "Oforikrom", 
        "Ayigya", "Bohyen", "Kaase", "Danyame", "Amakom", "Buokrom", "Kwadaso", "Nhyiaeso", 
        "Abuakwa", "Mampong", "Atonsu", "Ahinsan", "Boadi", "Kentinkrono", "Ayeduase", "Kotei", 
        "Santasi", "Patase", "Abrepo", "Bekwai", "Obuasi", "Konongo", "Effiduase", "Mamponteng"
    ],
    "Western": [
        "Takoradi Central", "Sekondi", "Effiakuma", "Anaji", "Kwesimintsim", "Apremdo", 
        "Kojokrom", "Nkontompo", "Fijai", "Nkroful", "Adiembra", "Essikado", "Tarkwa", 
        "Nsuta", "Tamso", "Prestea", "Bogoso"
    ],
    "Central": [
        "Cape Coast Central", "Pedu", "Adisadel", "Amamoma", "Kakumdo", "Kasoa Central", 
        "Opeikuma", "Akweley", "Budumburam", "Millennium City", "Winneba", "Elmina"
    ],
    "Eastern": [
        "Koforidua Central", "Adweso", "Effiduase", "Oyoko", "New Juaben", "Nsawam", 
        "Akwatia", "Asamankese", "Akim Oda", "Akim Swedru", "Anyinam", "Suhum", "Aburi"
    ],
    "Volta": [
        "Ho Central", "Hliha", "Ahoe", "Bankoe", "Deme", "Dome", "Klefe", "Hohoe", 
        "Keta", "Anloga", "Denu", "Sogakope", "Kpando"
    ],
    "Northern": [
        "Tamale Central", "Jisonayili", "Dungu", "Lamashegu", "Tishigu", "Yendi", 
        "Bimbilla", "Gushegu", "Karaga", "Savelugu"
    ],
    "Upper East": ["Bolgatanga Central", "Zuarungu", "Bongo", "Navrongo", "Paga", "Bawku"],
    "Upper West": ["Wa Central", "Dafiama", "Gwolu", "Jirapa", "Lawra", "Nandom"],
    "Western North": ["Sefwi Wiawso", "Bibiani", "Sefwi Bekwai", "Akontombra"],
    "Savannah": ["Damango", "Bole", "Salaga", "Daboya", "Yapei"],
    "North East": ["Nalerigu", "Gambaga", "Walewale", "Bunkpurugu"],
    "Oti": ["Dambai", "Jasikan", "Kadjebi", "Kete Krachi", "Nkwanta"],
    "Ahafo": ["Goaso", "Bechem", "Duayaw Nkwanta", "Kenyasi", "Mim"],
    "Bono East": ["Techiman", "Kintampo", "Nkoranza", "Atebubu", "Jema"],
    "Bono": ["Sunyani Central", "Sunyani West", "Berekum", "Dormaa Ahenkro", "Nsoatre"]
}

# ============ SEARCH QUERIES ============
SEARCH_QUERIES = {
    "Hotels": "hotel",
    "SMEs": "small business company",
    "Restaurants": "restaurant",
    "Schools": "school international",
    "Hospitals": "hospital medical center",
    "Shopping Malls": "shopping mall",
    "Supermarkets": "supermarket",
    "Pharmaceuticals": "pharmacy drug store",
    "Insurance": "insurance company",
    "Law Firms": "law firm solicitor",
    "Tech Companies": "technology IT company",
    "Real Estate": "real estate agency",
    "Construction": "construction company",
    "Logistics": "logistics shipping company",
    "NGOs": "non governmental organization",
    "Printing": "printing services",
    "Event Planning": "event planning services"
}

# ============ QR CODE GENERATION ============
def generate_qr_code_base64(url):
    try:
        qr = qrcode.QRCode(version=1, box_size=10, border=4)
        qr.add_data(url)
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="black", back_color="white")
        buffered = BytesIO()
        qr_img.save(buffered, format="PNG")
        return base64.b64encode(buffered.getvalue()).decode()
    except:
        return ""

def create_docx_letter(company_name, company_phone="", company_address="", qr_base64=""):
    """Create DOCX letter with QR code"""
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    import base64
    from PIL import Image as PILImage
    
    doc = Document()
    
    # Header
    header = doc.add_heading('TECHWOKX IT SOLUTIONS', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('FREE 5-Step Email & IT Health Check', 1)
    
    # Date
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()
    
    # Address
    doc.add_paragraph(f"To: Management Team")
    doc.add_paragraph(f"{company_name}")
    if company_address:
        doc.add_paragraph(f"{company_address}")
    doc.add_paragraph("Ghana")
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("Dear Management Team,")
    doc.add_paragraph()
    
    # Introduction
    doc.add_paragraph("I hope this letter finds you well.")
    doc.add_paragraph("My name is George Jabley, and I represent TechWokx IT Solutions. We help businesses improve their technology operations through reliable email systems, IT support, backup solutions, infrastructure reviews, and business continuity planning.")
    doc.add_paragraph()
    doc.add_paragraph("As part of our business outreach program, we conducted a high-level review of your organization's publicly available digital presence, including your website, domain configuration, SSL certificate status, and email setup.")
    doc.add_paragraph()
    doc.add_paragraph("Based on our observations, we identified several areas that may benefit from further review.")
    doc.add_paragraph()
    
    # What We Observed
    doc.add_heading('What We Observed', level=2)
    observations = [
        "• We were unable to verify whether all recommended email authentication records are fully configured.",
        "• Your company logo does not currently appear alongside emails in supported inboxes.",
        "• There may be opportunities to improve email deliverability and trust indicators.",
        "• We could not verify the existence of backup and recovery measures protecting business-critical communications and data."
    ]
    for obs in observations:
        doc.add_paragraph(obs)
    doc.add_paragraph()
    doc.add_paragraph("Please note that this review was conducted using publicly available information and from outside your network environment. As such, some findings may already have been addressed internally, and a more detailed assessment would be required to confirm their status.")
    doc.add_paragraph()
    
    # Free Offer
    doc.add_heading('Find out yourself: 5-Step Email & IT Health Check', level=2)
    doc.add_paragraph("To receive your complimentary assessment:")
    doc.add_paragraph("Scan the QR code on this letter, or")
    doc.add_paragraph("Visit: techwokx.online/#audit")
    doc.add_paragraph()
    doc.add_paragraph("The assessment takes less than five minutes and provides:")
    doc.add_paragraph("• A personalized technology health score")
    doc.add_paragraph("• Email configuration insights")
    doc.add_paragraph()
    
    # Why This Matters
    doc.add_heading('Why This Matters', level=2)
    doc.add_paragraph("Modern businesses rely heavily on email, cloud services, and digital communication. When these systems are not properly configured or maintained, organizations may experience:")
    issues = [
        "• Emails being delivered to spam or junk folders",
        "• Reduced trust in business communications",
        "• Difficulty identifying fraudulent communications",
        "• Loss of important business data",
        "• Increased downtime during technical issues",
        "• Challenges recovering from accidental deletions or system failures"
    ]
    for issue in issues:
        doc.add_paragraph(issue)
    doc.add_paragraph()
    
    # How TechWokx Can Help
    doc.add_heading('How TechWokx Can Help', level=2)
    doc.add_paragraph("We provide practical technology solutions designed to improve reliability, productivity, and business continuity.")
    doc.add_paragraph()
    
    # Services
    doc.add_heading('OUR SERVICES', level=2)
    services = [
        "⚡ Entry — Fast Fix: Email security setup, DNS configuration, professional signatures",
        "🔄 Recurring IT Retainer: 24/7 remote support, network troubleshooting, system health checks",
        "📋 Infrastructure Audit: Complete IT environment audit with risk report",
        "🤖 Automation Scripts: Custom automation for your business processes",
        "📧 FREE Professional HTML Email Signature Design - Branded signature with contact info"
    ]
    for service in services:
        doc.add_paragraph(service)
    doc.add_paragraph()
    
    # Backup & Business Continuity
    doc.add_heading('Backup & Business Continuity', level=2)
    doc.add_paragraph("We also assist organizations with secure backup and recovery solutions, helping businesses:")
    doc.add_paragraph("• Protect important files and emails")
    doc.add_paragraph("• Reduce downtime")
    doc.add_paragraph("• Recover quickly from accidental data loss")
    doc.add_paragraph("• Improve operational resilience")
    doc.add_paragraph()
    
    # Contact
    doc.add_heading('Let\'s Talk', level=2)
    doc.add_paragraph("If you would like to discuss your current IT environment, technology challenges, or future projects, we would be happy to schedule a brief conversation.")
    doc.add_paragraph()
    doc.add_paragraph("WhatsApp: +233 555 087 407")
    doc.add_paragraph("Email: hello@techwokx.online")
    doc.add_paragraph()
    doc.add_paragraph("At Techwokx, our approach is simple: Diagnose before we prescribe.")
    doc.add_paragraph()
    doc.add_paragraph("Thank you for your time and consideration. We look forward to the opportunity to support your organization.")
    doc.add_paragraph()
    
    # Signature
    doc.add_paragraph("Kind regards,")
    doc.add_paragraph()
    doc.add_paragraph("George Jabley")
    doc.add_paragraph("Founder & IT Operations Lead")
    doc.add_paragraph("TechWokx Ghana")
    
    # Add QR code
    if qr_base64:
        try:
            qr_data = base64.b64decode(qr_base64)
            qr_buffer = BytesIO(qr_data)
            doc.add_picture(qr_buffer, width=Inches(1.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    # Save to bytes
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# ============ WEBSITE & EMAIL FUNCTIONS ============
def check_website_status(url):
    if not url:
        return {"working": False, "error": "No website"}
    try:
        if not url.startswith("http"):
            url = "https://" + url
        response = requests.get(url, timeout=5, verify=False)
        return {"working": response.status_code == 200, "status_code": response.status_code}
    except:
        return {"working": False, "error": "Connection failed"}

def search_website_by_name(serp_api_key, business_name):
    if not serp_api_key:
        return None
    try:
        url = "https://serpapi.com/search"
        params = {
            "api_key": serp_api_key,
            "q": f"{business_name} Ghana official website",
            "engine": "google",
            "num": 5
        }
        response = requests.get(url, params=params, timeout=10)
        data = response.json()
        if "knowledge_graph" in data:
            website = data["knowledge_graph"].get("website", "")
            if website:
                return website
        if "organic_results" in data:
            for result in data["organic_results"][:3]:
                link = result.get("link", "")
                if link and "facebook" not in link and "twitter" not in link:
                    return link
        return None
    except:
        return None

def generate_possible_emails(company_name, domain=None):
    emails = []
    name_parts = company_name.lower().replace('&', 'and').split()
    base_name = ''.join(name_parts[:2]) if len(name_parts) > 1 else name_parts[0]
    base_name = re.sub(r'[^a-z0-9]', '', base_name)
    patterns = [
        f"info@{base_name}.com",
        f"contact@{base_name}.com",
        f"hello@{base_name}.com",
        f"support@{base_name}.com",
        f"admin@{base_name}.com",
        f"{base_name}@{base_name}.com"
    ]
    if domain:
        clean_domain = domain.replace('www.', '').replace('http://', '').replace('https://', '').split('/')[0]
        patterns = [f"info@{clean_domain}", f"contact@{clean_domain}", f"hello@{clean_domain}", f"admin@{clean_domain}"]
    return patterns

def enrich_company_data(company, serp_api_key):
    if company.get('website'):
        status = check_website_status(company['website'])
        company['website_status'] = status
    if not company.get('website') and serp_api_key:
        found_website = search_website_by_name(serp_api_key, company['name'])
        if found_website:
            company['website'] = found_website
            company['website_status'] = check_website_status(found_website)
    if not company.get('email'):
        company['possible_emails'] = generate_possible_emails(company['name'], company.get('website'))
    else:
        company['possible_emails'] = []
    return company

# ============ PERSONALIZED LETTER CONTENT ============
def get_personalized_letter(company_name, phone="", address=""):
    """Generate personalized letter content"""
    qr_base64 = generate_qr_code_base64("https://techwokx.online/#audit")
    current_date = datetime.now().strftime("%B %d, %Y")
    
    return f"""
================================================================================
                          TECHWOKX IT SOLUTIONS
                    FREE 5-Step Email & IT Health Check
================================================================================

Date: {current_date}

To: Management Team
{company_name}
{address}
Ghana

Dear Management Team,

I hope this letter finds you well.

My name is George Jabley, and I represent TechWokx IT Solutions. We help businesses improve their technology operations through reliable email systems, IT support, backup solutions, infrastructure reviews, and business continuity planning.

As part of our business outreach program, we conducted a high-level review of your organization's publicly available digital presence, including your website, domain configuration, SSL certificate status, and email setup.

Based on our observations, we identified several areas that may benefit from further review.

================================================================================
                          WHAT WE OBSERVED
================================================================================

• We were unable to verify whether all recommended email authentication records are fully configured.
• Your company logo does not currently appear alongside emails in supported inboxes.
• There may be opportunities to improve email deliverability and trust indicators.
• We could not verify the existence of backup and recovery measures protecting business-critical communications and data.

Please note that this review was conducted using publicly available information and from outside your network environment. 

As such, some findings may already have been addressed internally, and a more detailed assessment would be required to confirm their status.

================================================================================
                    FIND OUT YOURSELF: 5-STEP EMAIL & IT HEALTH CHECK
================================================================================

To receive your complimentary assessment:
• Scan the QR code on this letter
• Or visit: techwokx.online/#audit

The assessment takes less than five minutes and provides:
• A personalized technology health score
• Email configuration insights

================================================================================
                          WHY THIS MATTERS
================================================================================

Modern businesses rely heavily on email, cloud services, and digital communication. When these systems are not properly configured or maintained, organizations may experience:

• Emails being delivered to spam or junk folders
• Reduced trust in business communications
• Difficulty identifying fraudulent communications
• Loss of important business data
• Increased downtime during technical issues
• Challenges recovering from accidental deletions or system failures

================================================================================
                         HOW TECHWOKX CAN HELP
================================================================================

We provide practical technology solutions designed to improve reliability, productivity, and business continuity.

OUR SERVICES:

⚡ Entry — Fast Fix: Email security setup, DNS configuration, professional signatures

🔄 Recurring IT Retainer: 24/7 remote support, network troubleshooting, system health checks

📋 Infrastructure Audit: Complete IT environment audit with risk report and fix recommendations

🤖 Automation Scripts: Custom automation for your business processes

FREE COMPLIMENTARY OFFER:
As a courtesy, we would like to offer a FREE Professional HTML Email Signature Design for one member of your team at no cost.

The signature will include:
• Company branding
• Contact information
• Website and social media links
• Mobile-friendly design
• Outlook and Gmail compatibility

No obligation. No sales pressure. Just useful information to help you understand your current setup.

================================================================================
                          BACKUP & BUSINESS CONTINUITY
================================================================================

We also assist organizations with secure backup and recovery solutions through trusted technology partners, helping businesses:
• Protect important files and emails
• Reduce downtime
• Recover quickly from accidental data loss
• Improve operational resilience

================================================================================
                            LET'S TALK
================================================================================

If you would like to discuss your current IT environment, technology challenges, or future projects, we would be happy to schedule a brief conversation.

📞 WhatsApp: +233 555 087 407
✉️ Email: hello@techwokx.online
🌐 Website: techwokx.online

At Techwokx, our approach is simple: Diagnose before we prescribe.

Thank you for your time and consideration. We look forward to the opportunity to support your organization.

Kind regards,

George Jabley
Founder & IT Operations Lead
TechWokx Ghana

P.O. Box ML469, Malam, Accra, Ghana
================================================================================
"""

# ============ CALL SCRIPT ============
def generate_call_script(company_name, phone, bounce_reason=""):
    return f"""
================================================================================
                          CALL SCRIPT - TECHWOKX IT SOLUTIONS
================================================================================

Company: {company_name}
Phone: {phone}
Date: _______________

{"BOUNCE REASON: " + bounce_reason if bounce_reason else ""}

----------------- INTRODUCTION (15 seconds) -----------------
"Hello, this is George from TechWokx IT Solutions. 
Am I speaking with the person responsible for IT or business operations?"

----------------- IF YES -----------------
"Great! We recently attempted to send you information about our FREE IT assessment, 
but it appears the email may not have gone through."

"Let me briefly explain: We help businesses fix email and IT problems that 
are costing them time and money - things like emails going to spam, 
website issues, printer problems, and data backup."

----------------- THE OFFER -----------------
"We're offering a FREE 5-Step Email & IT Health Check that takes less than 5 minutes.
It will give you a personalized report showing exactly what needs attention."

----------------- QUESTIONS TO ASK -----------------
1. "Have you noticed any issues with your business email or internet?"
2. "Do you have a backup system for your important files?"
3. "Are you concerned about email security?"

----------------- NEXT STEPS -----------------
• If interested: "I'll resend the assessment link to your email. What's the best email address?"
• If not interested: "Would you prefer I mail you a physical letter with the information?"
• If busy: "What's a better time to reach you?"

----------------- CLOSING -----------------
"Thank you for your time. I'll follow up as promised."

================================================================================
NOTES:
Call Outcome: [Interested / Not Interested / Call Back / Wrong Number]
Follow-up Date: _______________
Email Collected: _______________
Proposal Requested: Yes / No
================================================================================
"""

# ============ GOOGLE PLACES API ============
def search_google_places(api_key, query, location):
    if not api_key:
        return []
    try:
        geocode_url = "https://maps.googleapis.com/maps/api/geocode/json"
        geocode_params = {"key": api_key, "address": f"{location}, Ghana"}
        geocode_response = requests.get(geocode_url, params=geocode_params, timeout=10)
        geocode_data = geocode_response.json()
        if not geocode_data.get('results'):
            return []
        lat = geocode_data['results'][0]['geometry']['location']['lat']
        lng = geocode_data['results'][0]['geometry']['location']['lng']
        places_url = "https://maps.googleapis.com/maps/api/place/nearbysearch/json"
        places_params = {
            "key": api_key, "location": f"{lat},{lng}", "radius": 5000,
            "keyword": query, "type": "establishment"
        }
        response = requests.get(places_url, params=places_params, timeout=10)
        data = response.json()
        businesses = []
        for place in data.get('results', [])[:20]:
            details_url = "https://maps.googleapis.com/maps/api/place/details/json"
            details_params = {
                "key": api_key, "place_id": place['place_id'],
                "fields": "name,formatted_address,formatted_phone_number,website,rating"
            }
            details_response = requests.get(details_url, params=details_params, timeout=10)
            details = details_response.json().get('result', {})
            businesses.append({
                "name": place.get('name'),
                "address": details.get('formatted_address', place.get('vicinity', '')),
                "phone": details.get('formatted_phone_number', ''),
                "website": details.get('website', ''),
                "email": "",
                "rating": details.get('rating', 0),
                "place_id": place.get('place_id')
            })
        return businesses
    except Exception as e:
        return []

def send_email(to_email, subject, body):
    if SMTP_CONFIG["sent_today"] >= SMTP_CONFIG["daily_limit"]:
        return False, "Daily limit reached"
    try:
        msg = MIMEMultipart()
        msg['From'] = f"TechWokx IT Solutions <{SMTP_CONFIG['username']}>"
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'html'))
        if SMTP_CONFIG['use_ssl']:
            server = smtplib.SMTP_SSL(SMTP_CONFIG['host'], SMTP_CONFIG['port'])
        else:
            server = smtplib.SMTP(SMTP_CONFIG['host'], SMTP_CONFIG['port'])
            server.starttls()
        server.login(SMTP_CONFIG['username'], SMTP_CONFIG['password'])
        server.send_message(msg)
        server.quit()
        SMTP_CONFIG["sent_today"] += 1
        return True, "Email sent"
    except Exception as e:
        return False, str(e)

def add_lead(lead_data):
    for lead in st.session_state.leads:
        if lead.get('name', '').lower() == lead_data.get('name', '').lower():
            return False
    new_lead = {
        "id": len(st.session_state.leads) + 1,
        "name": lead_data.get("name"),
        "address": lead_data.get("address", ""),
        "phone": lead_data.get("phone", ""),
        "website": lead_data.get("website", ""),
        "email": lead_data.get("email", ""),
        "possible_emails": lead_data.get("possible_emails", []),
        "website_status": lead_data.get("website_status", {}),
        "category": lead_data.get("category", ""),
        "rating": lead_data.get("rating", 0),
        "lead_score": 75,
        "created_at": datetime.now().isoformat(),
        "email_sent": False,
        "email_sent_date": None,
        "email_responded": False,
        "call_made": False,
        "call_notes": "",
        "followup_sequence": 0
    }
    st.session_state.leads.append(new_lead)
    save_leads()
    return True

def update_email_sent(lead_name, email):
    for lead in st.session_state.leads:
        if lead['name'] == lead_name:
            lead['email_sent'] = True
            lead['email_sent_date'] = datetime.now().isoformat()
            save_leads()
            return True
    return False

def update_email_responded(lead_name):
    for lead in st.session_state.leads:
        if lead['name'] == lead_name:
            lead['email_responded'] = True
            save_leads()
            return True
    return False

# ============ EMAIL BODY ============
def get_email_body(company_name):
    audit_url = "https://techwokx.online/#audit"
    return f"""
    <html>
    <body>
        <h2>FREE 5-Step Email & IT Health Check</h2>
        <p>Dear Management Team,</p>
        <p>My name is George Jabley from TechWokx IT Solutions.</p>
        <p>We conducted a review of your digital presence and identified opportunities to improve your IT infrastructure.</p>
        <p><strong>Get your Business Email Risk Score:</strong><br>
        <a href="{audit_url}">techwokx.online/#audit</a></p>
        <p>Best regards,<br>George Jabley<br>TechWokx Ghana</p>
    </body>
    </html>
    """

# ============ CSS ============
st.markdown("""
<style>
.stApp { background: #f8fafc; }
[data-testid="stSidebar"] { background: #0f172a; }
[data-testid="stSidebar"] .stMarkdown, [data-testid="stSidebar"] p { color: #e2e8f0; }
[data-testid="stSidebar"] .stButton button { background: rgba(251, 191, 36, 0.15); color: #fbbf24; width: 100%; margin: 2px 0; }
.welcome-card { background: linear-gradient(135deg, #667eea, #764ba2); border-radius: 16px; padding: 1.5rem; margin-bottom: 1.5rem; color: white; }
.metric-card { background: white; border-radius: 12px; padding: 1rem; border: 1px solid #e2e8f0; text-align: center; }
.metric-value { font-size: 1.8rem; font-weight: 700; color: #0f172a; }
.metric-label { color: #64748b; font-size: 0.8rem; }
.data-card { background: white; border-radius: 12px; padding: 1.2rem; border: 1px solid #e2e8f0; margin-bottom: 1rem; }
.section-header { color: #0f172a; font-size: 1.2rem; font-weight: 600; border-left: 3px solid #667eea; padding-left: 1rem; margin: 1rem 0; }
.custom-divider { height: 1px; background: #e2e8f0; margin: 1rem 0; }
.stButton > button { background: linear-gradient(135deg, #667eea, #764ba2); color: white; font-weight: 600; border: none; border-radius: 8px; }
</style>
""", unsafe_allow_html=True)

# ============ LOGIN ============
if not st.session_state.authenticated:
    st.markdown("""
    <div style="max-width: 400px; margin: 100px auto; padding: 2rem; background: white; border-radius: 24px; text-align: center;">
        <h1 style="color: #667eea;">TechWokx</h1>
        <p>Lead Intelligence Suite</p>
    </div>
    """, unsafe_allow_html=True)
    with st.form("login_form"):
        email = st.text_input("Email")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login")
        if submitted:
            if email in USERS and USERS[email]["password"] == password:
                st.session_state.authenticated = True
                st.session_state.user = USERS[email]
                st.session_state.current_page = 'dashboard'
                st.rerun()
            else:
                st.error("Invalid credentials")
    st.stop()

# ============ SIDEBAR ============
with st.sidebar:
    st.markdown("### TechWokx")
    st.markdown(f"Welcome, {st.session_state.user['name']}")
    st.markdown("---")
    menu_options = [
        ("🏠 Dashboard", "dashboard"),
        ("🔍 Search Companies", "search"),
        ("👥 Leads CRM", "crm"),
        ("📞 Call Queue", "call_queue"),
        ("🏢 Visit Queue", "visit_queue"),
        ("📧 Email Log", "email_log"),
        ("⚙️ Settings", "settings"),
        ("🚪 Logout", "logout")
    ]
    for label, page in menu_options:
        if st.button(label, key=f"nav_{page}", use_container_width=True):
            if page == "logout":
                st.session_state.authenticated = False
                st.rerun()
            else:
                st.session_state.current_page = page
                st.rerun()
    st.markdown("---")
    api_keys = get_api_keys()
    st.markdown(f"Google Maps: {'✅' if api_keys['google_maps'] else '❌'}")
    st.markdown(f"SERP API: {'✅' if api_keys['serp_api'] else '❌'}")
    st.markdown(f"Leads: {len(st.session_state.leads)}")
    st.markdown(f"Emails Today: {SMTP_CONFIG['sent_today']}/{SMTP_CONFIG['daily_limit']}")

# ============ DASHBOARD ============
if st.session_state.current_page == 'dashboard':
    st.markdown("""
    <div class="welcome-card">
        <h2>TechWokx Lead Intelligence</h2>
        <p>Professional IT Outreach | Email Tracking | Call Management | Visit Planning</p>
    </div>
    """, unsafe_allow_html=True)
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{len(st.session_state.leads)}</div><div class='metric-label'>Total Leads</div></div>", unsafe_allow_html=True)
    with col2:
        emails_sent = len([l for l in st.session_state.leads if l.get('email_sent')])
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{emails_sent}</div><div class='metric-label'>Emails Sent</div></div>", unsafe_allow_html=True)
    with col3:
        responded = len([l for l in st.session_state.leads if l.get('email_responded')])
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{responded}</div><div class='metric-label'>Responded</div></div>", unsafe_allow_html=True)
    with col4:
        remaining = SMTP_CONFIG['daily_limit'] - SMTP_CONFIG['sent_today']
        st.markdown(f"<div class='metric-card'><div class='metric-value'>{remaining}</div><div class='metric-label'>Emails Left</div></div>", unsafe_allow_html=True)
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        st.markdown('<div class="section-header">📍 Quick Actions</div>', unsafe_allow_html=True)
        if st.button("🔍 Search Companies", use_container_width=True):
            st.session_state.current_page = 'search'
            st.rerun()
        if st.button("👥 View CRM", use_container_width=True):
            st.session_state.current_page = 'crm'
            st.rerun()
    with col2:
        st.markdown('<div class="section-header">📊 Recent Activity</div>', unsafe_allow_html=True)
        for lead in st.session_state.leads[-3:]:
            st.markdown(f"<div class='data-card'><strong>{lead['name']}</strong><br>Added: {lead['created_at'][:10]}</div>", unsafe_allow_html=True)

# ============ SEARCH COMPANIES ============
if st.session_state.current_page == 'search':
    st.markdown('<div class="section-header">🔍 Search Companies</div>', unsafe_allow_html=True)
    st.caption("Find businesses across Ghana | Real Google Places Data")
    st.markdown("---")
    api_keys = get_api_keys()
    if not api_keys["google_maps"]:
        st.error("❌ Google Maps API key not configured!")
        st.stop()
    
    col1, col2 = st.columns(2)
    with col1:
        regions = list(ALL_CITIES.keys())
        selected_region = st.selectbox("Select Region", regions)
    with col2:
        if selected_region:
            cities = ALL_CITIES[selected_region]
            selected_city = st.selectbox("Select City/Town", cities)
    
    col1, col2 = st.columns(2)
    with col1:
        categories = list(SEARCH_QUERIES.keys())
        selected_category = st.selectbox("Business Category", categories)
    with col2:
        limit = st.slider("Number of Companies", 5, 30, 15)
    
    if st.button("🔍 Search", type="primary"):
        if selected_city:
            with st.spinner(f"Searching for {selected_category} in {selected_city}..."):
                query = SEARCH_QUERIES.get(selected_category, selected_category)
                businesses = search_google_places(api_keys["google_maps"], query, selected_city)
                for biz in businesses:
                    biz["category"] = selected_category
                    biz["city"] = selected_city
                    biz["region"] = selected_region
                st.session_state.batch_results = businesses[:limit]
                st.success(f"Found {len(businesses)} businesses")
    
    if st.session_state.batch_results:
        st.markdown("### Search Results")
        for idx, company in enumerate(st.session_state.batch_results):
            with st.expander(f"🏢 {company['name']}"):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**📍 Address:** {company.get('address', 'N/A')}")
                    st.markdown(f"**📞 Phone:** {company.get('phone', 'N/A')}")
                    website = company.get('website', '')
                    if website:
                        status = check_website_status(website)
                        st.markdown(f"**🌐 Website:** {website} {'✅' if status['working'] else '❌'}")
                    else:
                        st.markdown(f"**🌐 Website:** Not found")
                with col2:
                    st.markdown(f"**⭐ Rating:** {'⭐' * int(company.get('rating', 0))} {company.get('rating', 'N/A')}")
                st.markdown("---")
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    if st.button("📄 View Letter", key=f"letter_{idx}"):
                        letter_content = get_personalized_letter(company['name'], company.get('phone', ''), company.get('address', ''))
                        st.markdown(f"```\n{letter_content}\n```")
                with col2:
                    if st.button("➕ Add to CRM", key=f"add_{idx}"):
                        if add_lead(company):
                            st.success(f"Added {company['name']} to CRM")
                            st.rerun()
                        else:
                            st.warning("Already in CRM")
                with col3:
                    # Email input with custom option
                    email_options = ["[Enter Custom Email]"]
                    if company.get('email'):
                        email_options.append(company['email'])
                    possible_emails = generate_possible_emails(company['name'], company.get('website'))
                    email_options.extend(possible_emails[:3])
                    
                    selected_option = st.selectbox("Email", email_options, key=f"email_{idx}")
                    if selected_option == "[Enter Custom Email]":
                        custom_email = st.text_input("Enter Email", key=f"custom_{idx}")
                        final_email = custom_email
                    else:
                        final_email = selected_option
                    
                    if st.button("📧 Send Email", key=f"send_{idx}"):
                        if final_email:
                            email_body = get_email_body(company['name'])
                            subject = f"FREE IT Assessment for {company['name']}"
                            success, msg = send_email(final_email, subject, email_body)
                            if success:
                                update_email_sent(company['name'], final_email)
                                st.success(f"Email sent to {company['name']}")
                                st.rerun()
                            else:
                                st.error(f"Failed: {msg}")
                                # Add to call queue on failure
                                call_script = generate_call_script(company['name'], company.get('phone', ''), msg)
                                st.session_state.call_queue.append({
                                    "company": company['name'],
                                    "phone": company.get('phone', ''),
                                    "script": call_script,
                                    "reason": msg,
                                    "added_date": datetime.now().isoformat()
                                })
                                st.info("Added to call queue for follow-up")
                        else:
                            st.warning("Enter email")
                with col4:
                    if st.button("📞 Add to Call Queue", key=f"call_{idx}"):
                        call_script = generate_call_script(company['name'], company.get('phone', ''))
                        st.session_state.call_queue.append({
                            "company": company['name'],
                            "phone": company.get('phone', ''),
                            "script": call_script,
                            "added_date": datetime.now().isoformat()
                        })
                        st.success("Added to call queue")

# ============ CALL QUEUE ============
elif st.session_state.current_page == 'call_queue':
    st.markdown('<div class="section-header">📞 Call Queue</div>', unsafe_allow_html=True)
    st.caption("Companies to call - Use the call script")
    st.markdown("---")
    
    if st.session_state.call_queue:
        for idx, call in enumerate(st.session_state.call_queue):
            with st.expander(f"📞 {call['company']} - {call.get('phone', 'No phone')}"):
                st.markdown("### Call Script")
                st.code(call.get('script', 'No script'), language='markdown')
                if call.get('reason'):
                    st.warning(f"Reason: {call['reason']}")
                col1, col2 = st.columns(2)
                with col1:
                    outcome = st.selectbox("Call Outcome", ["Interested", "Not Interested", "Call Back", "Wrong Number", "No Answer"], key=f"outcome_{idx}")
                    if st.button("Save Result", key=f"save_{idx}"):
                        for lead in st.session_state.leads:
                            if lead['name'] == call['company']:
                                lead['call_made'] = True
                                lead['call_notes'] = f"Outcome: {outcome} on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
                                if outcome == "Interested" and lead.get('email'):
                                    lead['email_sent'] = False
                                save_leads()
                                st.success("Call result saved")
                                st.rerun()
                with col2:
                    if st.button("Remove", key=f"remove_{idx}"):
                        st.session_state.call_queue.pop(idx)
                        st.rerun()
    else:
        st.success("No pending calls")

# ============ VISIT QUEUE ============
elif st.session_state.current_page == 'visit_queue':
    st.markdown('<div class="section-header">🏢 Visit Queue</div>', unsafe_allow_html=True)
    st.markdown("---")
    if st.session_state.visit_queue:
        for idx, visit in enumerate(st.session_state.visit_queue):
            with st.expander(f"🏢 {visit['company']}"):
                st.markdown(f"**Address:** {visit.get('address', 'N/A')}")
                st.markdown(f"**Phone:** {visit.get('phone', 'N/A')}")
                if st.button("Mark Visited", key=f"visit_{idx}"):
                    st.session_state.visit_queue.pop(idx)
                    st.rerun()
    else:
        st.success("No pending visits")

# ============ LEADS CRM ============
elif st.session_state.current_page == 'crm':
    st.markdown('<div class="section-header">👥 Leads CRM</div>', unsafe_allow_html=True)
    st.caption(f"Total Leads: {len(st.session_state.leads)}")
    st.markdown("---")
    
    if st.session_state.leads:
        for lead in st.session_state.leads:
            with st.expander(f"🏢 {lead['name']} - Score: {lead['lead_score']}/100"):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**📍 Address:** {lead.get('address', 'N/A')}")
                    st.markdown(f"**📞 Phone:** {lead.get('phone', 'N/A')}")
                    st.markdown(f"**🌐 Website:** {lead.get('website', 'N/A')}")
                with col2:
                    st.markdown(f"**📧 Email:** {lead.get('email', 'N/A')}")
                    st.markdown(f"**📅 Added:** {lead['created_at'][:10]}")
                    st.markdown(f"**Email Sent:** {'✅' if lead.get('email_sent') else '❌'}")
                st.markdown("---")
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    if st.button(f"📄 Letter", key=f"letter_{lead['id']}"):
                        letter = get_personalized_letter(lead['name'], lead.get('phone', ''), lead.get('address', ''))
                        st.code(letter)
                with col2:
                    if st.button(f"📞 Add to Call", key=f"call_{lead['id']}"):
                        call_script = generate_call_script(lead['name'], lead.get('phone', ''))
                        st.session_state.call_queue.append({
                            "company": lead['name'],
                            "phone": lead.get('phone', ''),
                            "script": call_script,
                            "added_date": datetime.now().isoformat()
                        })
                        st.success("Added to call queue")
                with col3:
                    if st.button(f"🗑️ Delete", key=f"del_{lead['id']}"):
                        st.session_state.leads = [l for l in st.session_state.leads if l['id'] != lead['id']]
                        save_leads()
                        st.rerun()
                with col4:
                    if not lead.get('email_sent'):
                        email_options = ["[Enter Custom]"]
                        if lead.get('email'):
                            email_options.append(lead['email'])
                        if lead.get('possible_emails'):
                            email_options.extend(lead['possible_emails'][:2])
                        selected = st.selectbox("Email", email_options, key=f"email_crm_{lead['id']}")
                        if selected == "[Enter Custom]":
                            custom = st.text_input("Enter Email", key=f"custom_{lead['id']}")
                            final_email = custom
                        else:
                            final_email = selected
                        if st.button(f"📧 Send", key=f"send_{lead['id']}"):
                            if final_email:
                                email_body = get_email_body(lead['name'])
                                subject = f"FREE IT Assessment for {lead['name']}"
                                success, msg = send_email(final_email, subject, email_body)
                                if success:
                                    update_email_sent(lead['name'], final_email)
                                    st.success(f"Email sent")
                                    st.rerun()
                                else:
                                    st.error(f"Failed: {msg}")
    else:
        st.info("No leads yet")

# ============ EMAIL LOG ============
elif st.session_state.current_page == 'email_log':
    st.markdown('<div class="section-header">📧 Email Log</div>', unsafe_allow_html=True)
    st.markdown("---")
    if st.session_state.email_log:
        for log in reversed(st.session_state.email_log[-50:]):
            st.markdown(f"""
            <div class="data-card">
                <strong>To:</strong> {log['to']}<br>
                <strong>Company:</strong> {log['company']}<br>
                <strong>Date:</strong> {log['date'][:19]}<br>
                <strong>Status:</strong> {log['status']}
            </div>
            """, unsafe_allow_html=True)
    else:
        st.info("No emails sent yet")

# ============ SETTINGS ============
elif st.session_state.current_page == 'settings':
    st.markdown('<div class="section-header">⚙️ Settings</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("### API Configuration")
    st.code("""
    # .streamlit/secrets.toml
    GOOGLE_MAPS_API_KEY = "your_google_maps_api_key"
    SERP_API_KEY = "your_serp_api_key"
    """)
    st.markdown("### Email Configuration")
    st.info(f"""
    **SMTP Settings:**
    - Host: {SMTP_CONFIG['host']}
    - Port: {SMTP_CONFIG['port']}
    - Username: {SMTP_CONFIG['username']}
    - Daily Limit: {SMTP_CONFIG['daily_limit']} emails/day
    """)
    new_limit = st.number_input("Daily Email Limit", min_value=10, max_value=200, value=SMTP_CONFIG['daily_limit'])
    if st.button("Update Limit"):
        SMTP_CONFIG['daily_limit'] = new_limit
        st.success(f"Daily limit updated to {new_limit}")
    st.markdown("### Data Management")
    if st.button("Clear All Data", type="secondary"):
        st.session_state.leads = []
        st.session_state.email_log = []
        st.session_state.batch_results = []
        st.session_state.call_queue = []
        st.session_state.visit_queue = []
        save_leads()
        save_email_log()
        st.success("All data cleared")
        st.rerun()

# ============ FOOTER ============
st.markdown('<div class="custom-divider"></div>', unsafe_allow_html.html
